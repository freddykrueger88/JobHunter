"""DOCX-Vorlagen-Service: Platzhalter extrahieren, KI-befüllen, DOCX generieren.

Issue #89 – Anschreiben-Vorlage hochladen & KI-gestützt befüllen.

Unterstützte Platzhalter:
    {{FIRMA_NAME}}         – aus Stellendaten
    {{FIRMA_STRASSE}}      – aus Stellendaten (address)
    {{FIRMA_PLZ_ORT}}      – aus Stellendaten (postal_code + city)
    {{DATUM}}              – automatisch aktuelles Tagesdatum
    {{STELLE_BEZEICHNUNG}} – aus Stellendaten (title)
    {{ANSCHREIBEN_TEXT}}   – KI generiert individuellen Text
    {{ANSPRECHPARTNER}}    – aus Stellendaten (contact_person)
    {{BEWERBER_NAME}}      – aus CV-Daten
"""

import os
import re
import copy
import shutil
import tempfile
from datetime import date
from docx import Document as DocxDocument

import httpx
from backend.core.config import settings

# Regex für {{PLATZHALTER}}-Syntax
PLACEHOLDER_PATTERN = re.compile(r"\{\{([A-Z_]+)\}\}")

# Bekannte Platzhalter mit Beschreibungen
KNOWN_PLACEHOLDERS = {
    "FIRMA_NAME": "Firmenname aus Stellenanzeige",
    "FIRMA_STRASSE": "Straße der Firma",
    "FIRMA_PLZ_ORT": "PLZ und Ort der Firma",
    "DATUM": "Aktuelles Tagesdatum (automatisch)",
    "STELLE_BEZEICHNUNG": "Stellentitel",
    "ANSCHREIBEN_TEXT": "Von KI generierter Anschreiben-Text",
    "ANSPRECHPARTNER": "Ansprechpartner / Kontaktperson",
    "BEWERBER_NAME": "Name des Bewerbers (aus CV)",
}


def extract_placeholders(filepath: str) -> list[str]:
    """Liest ein DOCX und extrahiert alle {{PLATZHALTER}}-Markierungen.

    Sucht in Absätzen, Tabellen und Kopf-/Fußzeilen.
    """
    doc = DocxDocument(filepath)
    found: set[str] = set()

    # Absätze
    for para in doc.paragraphs:
        for match in PLACEHOLDER_PATTERN.finditer(para.text):
            found.add(match.group(1))

    # Tabellen
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for match in PLACEHOLDER_PATTERN.finditer(para.text):
                        found.add(match.group(1))

    # Kopf-/Fußzeilen
    for section in doc.sections:
        for header_footer in [section.header, section.footer]:
            if header_footer is not None:
                for para in header_footer.paragraphs:
                    for match in PLACEHOLDER_PATTERN.finditer(para.text):
                        found.add(match.group(1))

    return sorted(found)


def _replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    """Ersetzt Platzhalter in einem Absatz und behält Formatierung bei.

    python-docx speichert Text in Runs. Wenn ein Platzhalter über mehrere Runs
    verteilt ist (z.B. durch Word-Formatierung), werden die Runs zusammengeführt
    und der Text ersetzt.
    """
    full_text = paragraph.text
    if not PLACEHOLDER_PATTERN.search(full_text):
        return

    # Prüfe ob Ersetzungen nötig sind
    new_text = full_text
    for key, value in replacements.items():
        placeholder = "{{" + key + "}}"
        new_text = new_text.replace(placeholder, value)

    if new_text == full_text:
        return

    # Strategie: Alle Runs löschen, neuen Text in den ersten Run schreiben
    # Format des ersten Runs beibehalten
    if not paragraph.runs:
        return

    first_run = paragraph.runs[0]
    first_run.text = new_text

    # Restliche Runs leeren
    for run in paragraph.runs[1:]:
        run.text = ""


def fill_template(
    template_path: str,
    replacements: dict[str, str],
    output_path: str | None = None,
) -> str:
    """Befüllt eine DOCX-Vorlage mit den gegebenen Platzhalter-Werten.

    Args:
        template_path: Pfad zur Original-DOCX-Vorlage.
        replacements: Dict mit Platzhalter-Name → Wert (ohne {{ }}).
        output_path: Optionaler Ausgabepfad. Wenn None, wird ein temp-Pfad generiert.

    Returns:
        Pfad zur generierten DOCX-Datei.
    """
    if output_path is None:
        fd, output_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)

    # Kopiere Vorlage und arbeite auf der Kopie
    shutil.copy2(template_path, output_path)
    doc = DocxDocument(output_path)

    # Absätze
    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)

    # Tabellen
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, replacements)

    # Kopf-/Fußzeilen
    for section in doc.sections:
        for header_footer in [section.header, section.footer]:
            if header_footer is not None:
                for para in header_footer.paragraphs:
                    _replace_in_paragraph(para, replacements)

    doc.save(output_path)
    return output_path


def _format_date_german() -> str:
    """Gibt das aktuelle Datum im deutschen Format zurück, z.B. '2. Juni 2026'."""
    months_de = [
        "", "Januar", "Februar", "März", "April", "Mai", "Juni",
        "Juli", "August", "September", "Oktober", "November", "Dezember",
    ]
    today = date.today()
    return f"{today.day}. {months_de[today.month]} {today.year}"


def build_replacements_from_job(
    job: dict,
    cv_name: str | None = None,
    ai_text: str | None = None,
) -> dict[str, str]:
    """Baut das Replacements-Dict aus Job-Daten, CV und KI-Text.

    Args:
        job: Dict mit Job-Feldern (title, company, address, city, postal_code, contact_person).
        cv_name: Name des Bewerbers aus CV-Daten.
        ai_text: Von der KI generierter Anschreiben-Text.

    Returns:
        Dict mit Platzhalter-Name → Wert.
    """
    plz_ort_parts = []
    if job.get("postal_code"):
        plz_ort_parts.append(job["postal_code"])
    if job.get("city"):
        plz_ort_parts.append(job["city"])

    return {
        "FIRMA_NAME": job.get("company", ""),
        "FIRMA_STRASSE": job.get("address", "") or "",
        "FIRMA_PLZ_ORT": " ".join(plz_ort_parts),
        "DATUM": _format_date_german(),
        "STELLE_BEZEICHNUNG": job.get("title", ""),
        "ANSCHREIBEN_TEXT": ai_text or "",
        "ANSPRECHPARTNER": job.get("contact_person", "") or "Sehr geehrte Damen und Herren",
        "BEWERBER_NAME": cv_name or "",
    }


async def generate_cover_letter_text(
    job_title: str,
    company: str,
    contact_person: str | None,
    job_description: str | None,
    cv_summary: str | None,
    tone: str = "formell",
    model: str = "mistral",
    profile_summary: str | None = None,
) -> str:
    """Generiert den Anschreiben-Fließtext via Ollama.

    Gibt NUR den Textbody zurück – kein Datum, keine Adresse, kein Betreff.
    Diese Teile kommen aus den Platzhaltern der Vorlage.
    """
    tone_prompts = {
        "formell": "Schreibe in einem professionellen, formellen Stil. Siez-Form. Klare Struktur.",
        "direkt": "Schreibe direkt und auf den Punkt. Keine Füllwörter. Selbstbewusst.",
        "modern": "Schreibe modern und zeitgemäß. Du-Form ist ok. Frisch und authentisch.",
        "kreativ": "Schreibe kreativ und einprägsam. Hebe dich von der Masse ab. Originell.",
    }
    tone_instruction = tone_prompts.get(tone, tone_prompts["formell"])

    contact_line = f"Ansprechpartner: {contact_person}" if contact_person else "Kein Ansprechpartner bekannt"

    prompt = f"""Du bist ein Experte für Bewerbungsschreiben.
{tone_instruction}

Schreibe NUR den Fließtext eines Anschreibens (Anrede, 3 Absätze, Grußformel).
KEIN Datum, KEINE Adresse, KEIN Betreff – nur der eigentliche Text.
Gib NUR den fertigen Text zurück, keine Erklärungen.

Stelleninformationen:
- Position: {job_title}
- Unternehmen: {company}
- {contact_line}
- Stellenbeschreibung: {(job_description or 'Keine Beschreibung')[:1500]}

Bewerber-Profil:
{cv_summary or 'Kein CV vorhanden – schreibe einen allgemeinen Text.'}
{(chr(10) + chr(10) + profile_summary) if profile_summary else ''}
"""

    try:
        async with httpx.AsyncClient() as client:
            r = await client.post(
                f"{settings.OLLAMA_BASE_URL}/api/generate",
                json={"model": model, "prompt": prompt, "stream": False},
                timeout=180,
            )
            r.raise_for_status()
            return r.json().get("response", "").strip()
    except Exception as e:
        return f"Fehler bei KI-Generierung: {e}"
